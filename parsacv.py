import streamlit as st
import pandas as pd
import os
import re
from datetime import datetime, timedelta
from typing import List, Dict, Any, Optional
import json
from pathlib import Path
from dateutil import parser
from dateutil.relativedelta import relativedelta
import math

# Core libraries
import PyMuPDF  # fitz
import docx
from PIL import Image
import pytesseract
from langdetect import detect
import spacy
from spacy.matcher import Matcher, PhraseMatcher
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill

# Semantic similarity
from sentence_transformers import SentenceTransformer
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity

class AdvancedCVAnalyzer:
    def __init__(self):
        self.supported_languages = ['en', 'ar']
        self.nlp_models = {}
        self.matchers = {}
        self.phrase_matchers = {}
        self.semantic_model = None
        self.load_models()
        self.setup_patterns()
        
    def load_models(self):
        """Load all required models"""
        # Load spaCy models
        try:
            self.nlp_models['en'] = spacy.load("en_core_web_sm")
            st.success("✅ English NLP model loaded")
        except OSError:
            st.error("❌ English spaCy model not found. Install: python -m spacy download en_core_web_sm")
            
        try:
            self.nlp_models['ar'] = spacy.load("ar_core_news_sm")
            st.success("✅ Arabic NLP model loaded")
        except OSError:
            st.error("❌ Arabic spaCy model not found. Install: python -m spacy download ar_core_news_sm")
        
        # Load semantic similarity model
        try:
            self.semantic_model = SentenceTransformer('all-MiniLM-L6-v2')
            st.success("✅ Semantic similarity model loaded")
        except Exception as e:
            st.error(f"❌ Failed to load semantic model: {e}")

    def setup_patterns(self):
        """Setup advanced patterns for information extraction"""
        for lang_code, nlp in self.nlp_models.items():
            # Initialize matchers
            self.matchers[lang_code] = Matcher(nlp.vocab)
            self.phrase_matchers[lang_code] = PhraseMatcher(nlp.vocab, attr="LOWER")
            
            if lang_code == 'en':
                self.setup_english_patterns(nlp, lang_code)
            else:
                self.setup_arabic_patterns(nlp, lang_code)

    def setup_english_patterns(self, nlp, lang_code):
        """Setup English extraction patterns"""
        matcher = self.matchers[lang_code]
        phrase_matcher = self.phrase_matchers[lang_code]
        
        # Experience section headers
        experience_headers = [
            "professional experience", "work experience", "employment history",
            "career history", "work history", "professional background",
            "employment", "experience", "career", "work"
        ]
        
        # Education section headers
        education_headers = [
            "education", "educational background", "academic background",
            "qualifications", "academic qualifications", "degrees"
        ]
        
        # Job title patterns
        job_title_patterns = [
            [{"LOWER": {"IN": ["senior", "junior", "lead", "chief", "head", "principal"]}}, 
             {"POS": "NOUN", "OP": "+"}],
            [{"POS": "NOUN"}, {"LOWER": {"IN": ["engineer", "developer", "manager", "analyst", "specialist"]}}],
            [{"LOWER": {"IN": ["software", "web", "mobile", "data", "systems"]}}, 
             {"LOWER": {"IN": ["engineer", "developer", "architect"]}}]
        ]
        
        # Date patterns (improved)
        date_patterns = [
            [{"SHAPE": "dd/dd/dddd"}],  # 12/01/2020
            [{"SHAPE": "dd-dd-dddd"}],  # 12-01-2020
            [{"LIKE_NUM": True}, {"LOWER": {"IN": ["jan", "feb", "mar", "apr", "may", "jun",
                                                  "jul", "aug", "sep", "oct", "nov", "dec"]}}, 
             {"LIKE_NUM": True}],  # 15 Jan 2020
            [{"LOWER": {"IN": ["january", "february", "march", "april", "may", "june",
                              "july", "august", "september", "october", "november", "december"]}}, 
             {"LIKE_NUM": True}],  # January 2020
        ]
        
        # Company patterns
        company_patterns = [
            [{"ENT_TYPE": "ORG"}],
            [{"POS": "PROPN", "OP": "+"}, {"LOWER": {"IN": ["inc", "ltd", "llc", "corp", "corporation", "company"]}}],
            [{"IS_TITLE": True, "OP": "+"}, {"LOWER": {"IN": ["technologies", "systems", "solutions", "services"]}}]
        ]
        
        # Add patterns to matcher
        for i, pattern in enumerate(job_title_patterns):
            matcher.add(f"JOB_TITLE_{i}", [pattern])
        
        for i, pattern in enumerate(date_patterns):
            matcher.add(f"DATE_PATTERN_{i}", [pattern])
            
        for i, pattern in enumerate(company_patterns):
            matcher.add(f"COMPANY_{i}", [pattern])
        
        # Add phrase patterns
        experience_docs = [nlp(text) for text in experience_headers]
        education_docs = [nlp(text) for text in education_headers]
        
        phrase_matcher.add("EXPERIENCE_HEADER", experience_docs)
        phrase_matcher.add("EDUCATION_HEADER", education_docs)

    def setup_arabic_patterns(self, nlp, lang_code):
        """Setup Arabic extraction patterns"""
        matcher = self.matchers[lang_code]
        phrase_matcher = self.phrase_matchers[lang_code]
        
        # Arabic experience headers
        experience_headers_ar = [
            "الخبرة المهنية", "الخبرة العملية", "تاريخ العمل", "الخبرات",
            "المسيرة المهنية", "الوظائف السابقة", "خبرة", "عمل"
        ]
        
        # Arabic education headers
        education_headers_ar = [
            "التعليم", "المؤهلات الأكاديمية", "الخلفية التعليمية",
            "الشهادات", "التحصيل الأكاديمي", "الدراسة"
        ]
        
        # Add Arabic phrase patterns
        experience_docs_ar = [nlp(text) for text in experience_headers_ar]
        education_docs_ar = [nlp(text) for text in education_headers_ar]
        
        phrase_matcher.add("EXPERIENCE_HEADER", experience_docs_ar)
        phrase_matcher.add("EDUCATION_HEADER", education_docs_ar)

    def extract_text_from_pdf(self, file_path: str) -> str:
        """Enhanced PDF text extraction"""
        try:
            doc = PyMuPDF.open(file_path)
            text = ""
            for page in doc:
                # Try to extract text with layout preservation
                blocks = page.get_text("dict")
                page_text = ""
                
                for block in blocks["blocks"]:
                    if "lines" in block:
                        for line in block["lines"]:
                            line_text = ""
                            for span in line["spans"]:
                                line_text += span["text"]
                            page_text += line_text + "\n"
                
                text += page_text
            doc.close()
            return text
        except Exception as e:
            st.error(f"Error extracting PDF text: {e}")
            return ""

    def detect_language(self, text: str) -> str:
        """Enhanced language detection"""
        try:
            detected = detect(text)
            # More robust Arabic detection
            arabic_chars = re.findall(r'[\u0600-\u06FF]', text)
            english_chars = re.findall(r'[a-zA-Z]', text)
            
            if len(arabic_chars) > len(english_chars) * 0.3:
                return 'ar'
            return 'en' if detected in ['en', 'ar'] else 'en'
        except:
            return 'en'

    def extract_contact_info(self, text: str, language: str) -> Dict[str, str]:
        """Advanced contact information extraction"""
        contact_info = {
            'name': '',
            'email': '',
            'phone': '',
            'nationality': '',
            'age': ''
        }
        
        # Email extraction (improved)
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,7}\b'
        emails = re.findall(email_pattern, text, re.IGNORECASE)
        if emails:
            contact_info['email'] = emails[0]
        
        # Phone extraction (multiple international formats)
        phone_patterns = [
            r'\+?\d{1,4}[-.\s]?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,4}[-.\s]?\d{1,9}',
            r'\b\d{10,15}\b',
            r'\(\d{3,4}\)\s?\d{3,4}[-.]?\d{4}',  # (123) 456-7890
            r'\d{3,4}[-.]?\d{3,4}[-.]?\d{4}'     # 123-456-7890
        ]
        
        for pattern in phone_patterns:
            phones = re.findall(pattern, text)
            if phones:
                # Clean and validate phone number
                phone = re.sub(r'[^\d+]', '', phones[0])
                if 7 <= len(phone.replace('+', '')) <= 15:
                    contact_info['phone'] = phones[0]
                    break
        
        # Name extraction (advanced)
        lines = text.split('\n')[:15]  # Check first 15 lines
        
        for line in lines:
            line = line.strip()
            if not line or len(line) < 2:
                continue
                
            # Skip lines with common keywords
            skip_keywords = ['cv', 'resume', 'curriculum', 'vitae', 'phone', 'email', 'address']
            if any(keyword in line.lower() for keyword in skip_keywords):
                continue
            
            # Check if line could be a name
            words = line.split()
            if 1 <= len(words) <= 5:  # Names typically have 1-5 parts
                # Check if all words are likely names (no numbers, emails, etc.)
                if all(re.match(r'^[A-Za-z\u0600-\u06FF\s\'-]+$', word) for word in words):
                    if not any(char.isdigit() or '@' in line or 'www' in line.lower()):
                        contact_info['name'] = line
                        break
        
        # Age extraction
        age_patterns = [
            r'age:?\s*(\d{2})',
            r'العمر:?\s*(\d{2})',
            r'born:?\s*\d{4}',
            r'(\d{2})\s*years?\s*old'
        ]
        
        for pattern in age_patterns:
            age_match = re.search(pattern, text, re.IGNORECASE)
            if age_match:
                try:
                    age = int(age_match.group(1))
                    if 16 <= age <= 80:  # Reasonable age range
                        contact_info['age'] = str(age)
                        break
                except:
                    continue
        
        # Nationality extraction
        if language == 'en':
            nationality_pattern = r'nationality:?\s*([A-Za-z]+)'
            nat_match = re.search(nationality_pattern, text, re.IGNORECASE)
            if nat_match:
                contact_info['nationality'] = nat_match.group(1)
        
        return contact_info

    def parse_date(self, date_str: str) -> Optional[datetime]:
        """Advanced date parsing"""
        if not date_str:
            return None
        
        date_str = date_str.strip()
        
        # Handle "present", "current", etc.
        current_indicators = ['present', 'current', 'now', 'till date', 'الآن', 'حالياً']
        if any(indicator in date_str.lower() for indicator in current_indicators):
            return datetime.now()
        
        # Common date formats
        date_formats = [
            '%Y-%m-%d', '%d-%m-%Y', '%m-%d-%Y',
            '%Y/%m/%d', '%d/%m/%Y', '%m/%d/%Y',
            '%B %Y', '%b %Y', '%Y',
            '%m/%Y', '%Y-%m'
        ]
        
        for fmt in date_formats:
            try:
                return datetime.strptime(date_str, fmt)
            except:
                continue
        
        # Try using dateutil parser
        try:
            return parser.parse(date_str, fuzzy=True)
        except:
            return None

    def calculate_duration(self, start_date: datetime, end_date: datetime) -> Dict[str, Any]:
        """Calculate duration between dates"""
        if not start_date or not end_date:
            return {'years': 0, 'months': 0, 'total_months': 0, 'formatted': 'Unknown'}
        
        if end_date < start_date:
            start_date, end_date = end_date, start_date
        
        delta = relativedelta(end_date, start_date)
        total_months = delta.years * 12 + delta.months
        
        if delta.years > 0:
            formatted = f"{delta.years}y {delta.months}m"
        else:
            formatted = f"{delta.months}m"
        
        return {
            'years': delta.years,
            'months': delta.months,
            'total_months': total_months,
            'formatted': formatted
        }

    def extract_work_experience(self, text: str, language: str) -> List[Dict[str, Any]]:
        """Advanced work experience extraction using NLP patterns"""
        if language not in self.nlp_models:
            return []
        
        nlp = self.nlp_models[language]
        matcher = self.matchers[language]
        phrase_matcher = self.phrase_matchers[language]
        
        doc = nlp(text)
        experiences = []
        
        # Find experience section
        phrase_matches = phrase_matcher(doc)
        experience_start = 0
        education_start = len(doc)
        
        for match_id, start, end in phrase_matches:
            label = nlp.vocab.strings[match_id]
            if label == "EXPERIENCE_HEADER":
                experience_start = max(experience_start, start)
            elif label == "EDUCATION_HEADER":
                education_start = min(education_start, start)
        
        # Focus on experience section
        experience_doc = doc[experience_start:education_start]
        
        # Extract entities
        orgs = [ent.text.strip() for ent in experience_doc.ents if ent.label_ in ["ORG", "PERSON"]]
        dates = [ent.text.strip() for ent in experience_doc.ents if ent.label_ == "DATE"]
        gpes = [ent.text.strip() for ent in experience_doc.ents if ent.label_ == "GPE"]  # Countries/Cities
        
        # Use pattern matching to find job-related information
        matches = matcher(experience_doc)
        job_titles = []
        companies = []
        
        for match_id, start, end in matches:
            label = nlp.vocab.strings[match_id]
            matched_text = experience_doc[start:end].text.strip()
            
            if "JOB_TITLE" in label:
                job_titles.append(matched_text)
            elif "COMPANY" in label:
                companies.append(matched_text)
        
        # Combine extracted entities
        all_companies = list(set(orgs + companies))
        all_dates = dates
        
        # Pattern-based experience extraction
        lines = text.split('\n')
        current_job = None
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                if current_job and any(current_job.values()):
                    experiences.append(current_job)
                    current_job = None
                continue
            
            # Check if line contains company information
            line_companies = [comp for comp in all_companies if comp.lower() in line.lower()]
            line_dates = [date for date in all_dates if date in line]
            line_countries = [country for country in gpes if country in line]
            
            if line_companies or line_dates or any(keyword in line.lower() for keyword in 
                ['engineer', 'manager', 'developer', 'analyst', 'specialist', 'consultant',
                 'مهندس', 'مدير', 'مطور', 'محلل']):
                
                if current_job and any(current_job.values()):
                    experiences.append(current_job)
                
                current_job = {
                    'title': '',
                    'company': '',
                    'start_date': '',
                    'end_date': '',
                    'duration': {},
                    'country': '',
                    'description': ''
                }
                
                # Extract information from current line
                if line_companies:
                    current_job['company'] = line_companies[0]
                
                if line_countries:
                    current_job['country'] = line_countries[0]
                
                # Look for job title patterns
                for title in job_titles:
                    if title.lower() in line.lower():
                        current_job['title'] = title
                        break
                
                if not current_job['title']:
                    current_job['title'] = line
                
                # Extract dates from surrounding lines
                date_context = '\n'.join(lines[max(0, i-2):i+3])
                extracted_dates = re.findall(r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b|\b\d{4}\b|\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\s+\d{4}\b', date_context, re.IGNORECASE)
                
                if len(extracted_dates) >= 2:
                    start_date = self.parse_date(extracted_dates[0])
                    end_date = self.parse_date(extracted_dates[1])
                    
                    if start_date and end_date:
                        current_job['start_date'] = start_date.strftime('%Y-%m-%d')
                        current_job['end_date'] = end_date.strftime('%Y-%m-%d')
                        current_job['duration'] = self.calculate_duration(start_date, end_date)
                elif len(extracted_dates) == 1:
                    date_obj = self.parse_date(extracted_dates[0])
                    if date_obj:
                        current_job['start_date'] = date_obj.strftime('%Y-%m-%d')
                        current_job['end_date'] = datetime.now().strftime('%Y-%m-%d')
                        current_job['duration'] = self.calculate_duration(date_obj, datetime.now())
            
            elif current_job:
                # Add to description
                if current_job['description']:
                    current_job['description'] += ' ' + line
                else:
                    current_job['description'] = line
        
        # Add the last job if exists
        if current_job and any(current_job.values()):
            experiences.append(current_job)
        
        return experiences

    def extract_education(self, text: str, language: str) -> List[Dict[str, str]]:
        """Advanced education extraction"""
        if language not in self.nlp_models:
            return []
        
        nlp = self.nlp_models[language]
        phrase_matcher = self.phrase_matchers[language]
        
        doc = nlp(text)
        education = []
        
        # Find education section
        phrase_matches = phrase_matcher(doc)
        education_start = 0
        
        for match_id, start, end in phrase_matches:
            label = nlp.vocab.strings[match_id]
            if label == "EDUCATION_HEADER":
                education_start = max(education_start, start)
        
        # Focus on education section
        education_doc = doc[education_start:]
        
        # Extract educational entities
        orgs = [ent.text.strip() for ent in education_doc.ents if ent.label_ == "ORG"]
        dates = [ent.text.strip() for ent in education_doc.ents if ent.label_ == "DATE"]
        
        # Degree patterns
        degree_patterns = {
            'en': [
                r'\b(?:bachelor|master|phd|doctorate|diploma|certificate|degree)\b.*',
                r'\b(?:b\.?[as]\.?|m\.?[as]\.?|ph\.?d\.?|m\.?d\.?)\b.*',
                r'\b(?:undergraduate|graduate|postgraduate)\b.*'
            ],
            'ar': [
                r'بكالوريوس.*|ماجستير.*|دكتوراه.*|دبلوم.*|شهادة.*'
            ]
        }
        
        patterns = degree_patterns.get(language, degree_patterns['en'])
        
        lines = text.split('\n')
        current_edu = None
        
        for line in lines:
            line = line.strip()
            if not line:
                if current_edu and any(current_edu.values()):
                    education.append(current_edu)
                    current_edu = None
                continue
            
            # Check if line contains degree information
            is_degree_line = any(re.search(pattern, line, re.IGNORECASE) for pattern in patterns)
            line_institutions = [org for org in orgs if org.lower() in line.lower()]
            
            if is_degree_line or line_institutions:
                if current_edu and any(current_edu.values()):
                    education.append(current_edu)
                
                current_edu = {
                    'degree': '',
                    'institution': '',
                    'graduation_date': '',
                    'field': ''
                }
                
                if is_degree_line:
                    current_edu['degree'] = line
                
                if line_institutions:
                    current_edu['institution'] = line_institutions[0]
                
                # Extract graduation date
                line_dates = re.findall(r'\b\d{4}\b', line)
                if line_dates:
                    current_edu['graduation_date'] = line_dates[-1]  # Usually the latest year
            
            elif current_edu:
                # Add additional information
                if not current_edu['institution'] and any(org.lower() in line.lower() for org in orgs):
                    for org in orgs:
                        if org.lower() in line.lower():
                            current_edu['institution'] = org
                            break
        
        # Add the last education entry
        if current_edu and any(current_edu.values()):
            education.append(current_edu)
        
        return education

    def advanced_semantic_matching(self, cv_text: str, jd_text: str) -> Dict[str, Any]:
        """Advanced semantic matching using transformer models"""
        if not self.semantic_model:
            return self.basic_keyword_matching(cv_text, jd_text)
        
        try:
            # Split texts into sentences
            cv_sentences = [s.strip() for s in cv_text.split('.') if len(s.strip()) > 10]
            jd_sentences = [s.strip() for s in jd_text.split('.') if len(s.strip()) > 10]
            
            if not cv_sentences or not jd_sentences:
                return self.basic_keyword_matching(cv_text, jd_text)
            
            # Get embeddings
            cv_embeddings = self.semantic_model.encode(cv_sentences)
            jd_embeddings = self.semantic_model.encode(jd_sentences)
            
            # Calculate similarity matrix
            similarity_matrix = cosine_similarity(cv_embeddings, jd_embeddings)
            
            # Find best matches
            best_matches = []
            match_scores = []
            
            for jd_idx, jd_sentence in enumerate(jd_sentences):
                best_cv_idx = np.argmax(similarity_matrix[:, jd_idx])
                best_score = similarity_matrix[best_cv_idx, jd_idx]
                
                if best_score > 0.3:  # Threshold for meaningful similarity
                    best_matches.append({
                        'jd_sentence': jd_sentence[:100] + '...' if len(jd_sentence) > 100 else jd_sentence,
                        'cv_sentence': cv_sentences[best_cv_idx][:100] + '...' if len(cv_sentences[best_cv_idx]) > 100 else cv_sentences[best_cv_idx],
                        'score': float(best_score)
                    })
                    match_scores.append(best_score)
            
            # Calculate overall match score
            if match_scores:
                overall_score = (np.mean(match_scores) * 100)
            else:
                overall_score = 0
            
            # Extract skills using both approaches
            semantic_result = {
                'score': round(overall_score, 2),
                'best_matches': sorted(best_matches, key=lambda x: x['score'], reverse=True)[:5],
                'confidence': 'high' if len(match_scores) > 3 else 'medium'
            }
            
            # Combine with keyword matching for skill extraction
            keyword_result = self.basic_keyword_matching(cv_text, jd_text)
            
            return {
                'score': semantic_result['score'],
                'matched_skills': keyword_result['matched_skills'],
                'missing_skills': keyword_result['missing_skills'],
                'semantic_matches': semantic_result['best_matches'],
                'confidence': semantic_result['confidence']
            }
            
        except Exception as e:
            st.warning(f"Semantic matching failed, falling back to keyword matching: {e}")
            return self.basic_keyword_matching(cv_text, jd_text)

    def basic_keyword_matching(self, cv_text: str, jd_text: str) -> Dict[str, Any]:
        """Fallback keyword-based matching"""
        # Advanced preprocessing
        def preprocess_text(text):
            # Convert to lowercase
            text = text.lower()
            # Extract meaningful terms (2-4 words)
            terms = []
            words = re.findall(r'\b\w+\b', text)
            
            # Single important words
            important_single_words = {'python', 'java', 'javascript', 'react', 'angular', 'vue',
                                    'django', 'flask', 'spring', 'aws', 'azure', 'docker',
                                    'kubernetes', 'machine learning', 'ai', 'sql', 'nosql'}
            
            for word in words:
                if word in important_single_words:
                    terms.append(word)
            
            # Multi-word terms
            for i in range(len(words) - 1):
                bigram = f"{words[i]} {words[i+1]}"
                if len(bigram) > 6:  # Avoid short meaningless combinations
                    terms.append(bigram)
            
            for i in range(len(words) - 2):
                trigram = f"{words[i]} {words[i+1]} {words[i+2]}"
                if len(trigram) > 10:
                    terms.append(trigram)
            
            return set(terms)
        
        cv_terms = preprocess_text(cv_text)
        jd_terms = preprocess_text(jd_text)
        
        # Find matches
        matched_terms = cv_terms & jd_terms
        missing_terms = jd_terms - cv_terms
        
        # Calculate weighted score
        if jd_terms:
            exact_match_score = (len(matched_terms) / len(jd_terms)) * 100
            
            # Bonus for important technical terms
            tech_terms = {'python', 'java', 'machine learning', 'ai', 'aws', 'react'}
            tech_matches = sum(1 for term in matched_terms if any(tech in term for tech in tech_terms))
            tech_bonus = min(tech_matches * 5, 20)  # Max 20% bonus
            
            final_score = min(exact_match_score + tech_bonus, 100)
        else:
            final_score = 0
        
        return {
            'score': round(final_score, 2),
            'matched_skills': list(matched_terms)[:15],
            'missing_skills': list(missing_terms)[:15]
        }

    def calculate_total_experience(self, experiences: List[Dict]) -> Dict[str, Any]:
        """Calculate total work experience"""
        if not experiences:
            return {'years': 0, 'months': 0, 'total_months': 0}
        
        total_months = 0
        for exp in experiences:
            if exp.get('duration') and isinstance(exp['duration'], dict):
                total_months += exp['duration'].get('total_months', 0)
        
        years = total_months // 12
        months = total_months % 12
        
        return {
            'years': years,
            'months': months,
            'total_months': total_months
        }

    def process_cv(self, file_path: str, file_name: str, jd_text: str) -> Dict[str, Any]:
        """Process a single CV with advanced extraction"""
        file_extension = Path(file_path).suffix
        
        # Extract text
        if file_extension.lower() == '.pdf':
            cv_text = self.extract_text_from_pdf(file_path)
        elif file_extension.lower() == '.docx':
            cv_text = self.extract_text_from_docx(file_path)
        elif file_extension.lower() in ['.png', '.jpg', '.jpeg']:
            cv_text = self.extract_text_from_image(file_path)
        else:
            return None
        
        if not cv_text.strip():
            return None
        
        # Detect language
        language = self.detect_language(cv_text)
        
        # Extract information using advanced methods
        contact_info = self.extract_contact_info(cv_text, language)
        work_experience = self.extract_work_experience(cv_text, language)
        education = self.extract_education(cv_text, language)
        
        # Calculate total experience
        total_exp = self.calculate_total_experience(work_experience)
        
        # Calculate match score using semantic similarity
        match_result = self.advanced_semantic_matching(cv_text, jd_text)
        
        return {
            'file_name': file_name,
            'file_path': file_path,
            'language': language,
            'name': contact_info['name'] or 'Name not found',
            'email': contact_info['email'] or 'Email not found',
            'phone': contact_info['phone'] or 'Phone not found',
            'age': contact_info['age'] or 'Age not found',
            'nationality': contact_info['nationality'] or 'Nationality not found',
            'total_experience_years': total_exp['years'],
            'total_experience_months': total_exp['months'],
            'total_experience_formatted': f"{total_exp['years']}y {total_exp['months']}m" if total_exp['years'] > 0 else f"{total_exp['months']}m",
            'match_score': match_result['score'],
            'work_experience_raw': work_experience,
            'education_raw': education,
            'work_experience': self.format_experience_string(work_experience),
            'education': self.format_education_string(education),
            'matched_skills': ', '.join(match_result['matched_skills'][:10]),
            'missing_skills': ', '.join(match_result['missing_skills'][:10]),
            'semantic_matches': match_result.get('semantic_matches', []),
            'confidence': match_result.get('confidence', 'medium')
        }

    def extract_text_from_docx(self, file_path: str) -> str:
        """Enhanced DOCX text extraction"""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"
            
            return text
        except Exception as e:
            st.error(f"Error extracting DOCX text: {e}")
            return ""

    def extract_text_from_image(self, file_path: str) -> str:
        """Enhanced OCR text extraction"""
        try:
            image = Image.open(file_path)
            
            # Preprocess image for better OCR
            import cv2
            import numpy as np
            
            # Convert PIL to cv2
            opencv_image = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
            
            # Convert to grayscale
            gray = cv2.cvtColor(opencv_image, cv2.COLOR_BGR2GRAY)
            
            # Apply denoising
            denoised = cv2.fastNlMeansDenoising(gray)
            
            # Apply threshold
            _, thresh = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            
            # Convert back to PIL
            processed_image = Image.fromarray(thresh)
            
            # Configure tesseract for better results
            custom_config = r'--oem 3 --psm 6 -l eng+ara'
            text = pytesseract.image_to_string(processed_image, config=custom_config)
            
            # If OCR fails, try with original image
            if len(text.strip()) < 50:
                text = pytesseract.image_to_string(image, config=custom_config)
            
            return text
        except Exception as e:
            st.error(f"Error extracting text from image: {e}")
            return ""

    def format_experience_string(self, experience: List[Dict]) -> str:
        """Enhanced experience formatting"""
        if not experience:
            return "No work experience found"
        
        formatted = []
        for exp in experience:
            company = exp.get('company', 'Company not specified')
            title = exp.get('title', 'Position not specified')
            country = exp.get('country', 'Location not specified')
            
            if exp.get('duration') and isinstance(exp['duration'], dict):
                duration = exp['duration'].get('formatted', 'Duration unknown')
            else:
                duration = 'Duration unknown'
            
            formatted.append(f"[{title} at {company} | Location: {country} | Duration: {duration}]")
        
        return "; ".join(formatted)

    def format_education_string(self, education: List[Dict]) -> str:
        """Enhanced education formatting"""
        if not education:
            return "No education information found"
        
        formatted = []
        for edu in education:
            degree = edu.get('degree', 'Degree not specified')
            institution = edu.get('institution', 'Institution not specified')
            year = edu.get('graduation_date', 'Year not specified')
            
            formatted.append(f"[{degree} from {institution} ({year})]")
        
        return "; ".join(formatted)

    def generate_excel_report(self, results: List[Dict], language: str = 'en') -> str:
        """Generate comprehensive Excel report"""
        # Enhanced column headers
        headers = {
            'en': {
                'name': 'Full Name',
                'email': 'Email Address',
                'phone': 'Phone Number',
                'age': 'Age',
                'nationality': 'Nationality',
                'total_experience': 'Total Experience',
                'match_score': 'Match Score (%)',
                'confidence': 'Analysis Confidence',
                'work_experience': 'Work Experience Details',
                'education': 'Education Background',
                'matched_skills': 'Matched Skills',
                'missing_skills': 'Missing Skills',
                'language_detected': 'CV Language',
                'file_path': 'CV File Path'
            },
            'ar': {
                'name': 'الاسم الكامل',
                'email': 'عنوان البريد الإلكتروني',
                'phone': 'رقم الهاتف',
                'age': 'العمر',
                'nationality': 'الجنسية',
                'total_experience': 'إجمالي الخبرة',
                'match_score': 'نسبة التطابق (%)',
                'confidence': 'مستوى الثقة في التحليل',
                'work_experience': 'تفاصيل الخبرة العملية',
                'education': 'الخلفية التعليمية',
                'matched_skills': 'المهارات المطابقة',
                'missing_skills': 'المهارات المفقودة',
                'language_detected': 'لغة السيرة الذاتية',
                'file_path': 'مسار ملف السيرة الذاتية'
            }
        }
        
        # Create DataFrame with enhanced data
        df_data = []
        for result in results:
            if result:  # Skip None results
                df_data.append([
                    result['name'],
                    result['email'],
                    result['phone'],
                    result['age'],
                    result['nationality'],
                    result['total_experience_formatted'],
                    result['match_score'],
                    result.get('confidence', 'medium').title(),
                    result['work_experience'],
                    result['education'],
                    result['matched_skills'],
                    result['missing_skills'],
                    result['language'].upper(),
                    result['file_path']
                ])
        
        # Create DataFrame
        selected_headers = headers[language] if language in headers else headers['en']
        df = pd.DataFrame(df_data, columns=list(selected_headers.values()))
        
        # Sort by match score (descending)
        df = df.sort_values(by=list(selected_headers.values())[6], ascending=False)  # Match score column
        
        # Generate filename with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_file = f"ParsaCV_Analysis_Report_{timestamp}.xlsx"
        
        with pd.ExcelWriter(output_file, engine='openpyxl') as writer:
            # Main analysis sheet
            df.to_excel(writer, sheet_name='CV Analysis', index=False)
            
            # Create summary sheet
            summary_data = {
                'Metric': [
                    'Total CVs Analyzed',
                    'Average Match Score',
                    'Highest Match Score',
                    'Lowest Match Score',
                    'CVs Above 70% Match',
                    'CVs Above 50% Match',
                    'English CVs',
                    'Arabic CVs'
                ],
                'Value': [
                    len(df),
                    f"{df.iloc[:, 6].mean():.1f}%" if len(df) > 0 else "0%",
                    f"{df.iloc[:, 6].max():.1f}%" if len(df) > 0 else "0%",
                    f"{df.iloc[:, 6].min():.1f}%" if len(df) > 0 else "0%",
                    len(df[df.iloc[:, 6] >= 70]),
                    len(df[df.iloc[:, 6] >= 50]),
                    len(df[df.iloc[:, 12] == 'EN']),
                    len(df[df.iloc[:, 12] == 'AR'])
                ]
            }
            
            summary_df = pd.DataFrame(summary_data)
            summary_df.to_excel(writer, sheet_name='Summary', index=False)
            
            # Format both sheets
            for sheet_name in ['CV Analysis', 'Summary']:
                worksheet = writer.sheets[sheet_name]
                
                # Style headers
                header_font = Font(bold=True, color="FFFFFF", size=12)
                header_fill = PatternFill(start_color="2F5597", end_color="2F5597", fill_type="solid")
                
                for col in range(1, worksheet.max_column + 1):
                    cell = worksheet.cell(row=1, column=col)
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = Alignment(horizontal="center", vertical="center")
                
                # Auto-adjust column widths
                for column in worksheet.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    adjusted_width = min(max_length + 3, 60)
                    worksheet.column_dimensions[column_letter].width = adjusted_width
                
                # Add conditional formatting for match scores (only for main sheet)
                if sheet_name == 'CV Analysis' and len(df) > 0:
                    from openpyxl.formatting.rule import ColorScaleRule
                    
                    # Find match score column (7th column, index 6)
                    match_score_col = worksheet['G']  # Assuming match score is in column G
                    
                    color_scale = ColorScaleRule(
                        start_type='min', start_color='FFE6E6',  # Light red
                        mid_type='percentile', mid_value=50, mid_color='FFFFCC',  # Light yellow
                        end_type='max', end_color='E6FFE6'  # Light green
                    )
                    
                    worksheet.conditional_formatting.add(
                        f'G2:G{len(df) + 1}',
                        color_scale
                    )
        
        return output_file


def main():
    st.set_page_config(
        page_title="ParsaCV - Advanced Multilingual CV Analyzer",
        page_icon="🎯",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    # Custom CSS for better styling
    st.markdown("""
    <style>
    .main-header {
        text-align: center;
        padding: 2rem 0;
        background: linear-gradient(90deg, #2F5597 0%, #1E3A8A 100%);
        color: white;
        border-radius: 10px;
        margin-bottom: 2rem;
    }
    .metric-card {
        background: white;
        padding: 1rem;
        border-radius: 8px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        border-left: 4px solid #2F5597;
    }
    .stProgress > div > div > div > div {
        background-color: #2F5597;
    }
    </style>
    """, unsafe_allow_html=True)
    
    # Language selection in sidebar
    with st.sidebar:
        st.header("⚙️ Settings")
        language = st.selectbox(
            "🌐 Interface Language / لغة الواجهة",
            options=['en', 'ar'],
            format_func=lambda x: "🇺🇸 English" if x == 'en' else "🇸🇦 العربية"
        )
    
    # UI text based on selected language
    ui_text = {
        'en': {
            'title': "🎯 ParsaCV - Advanced Multilingual CV Analyzer",
            'subtitle': "Intelligent CV screening with advanced NLP and semantic matching",
            'features': [
                "🤖 Advanced NLP pattern matching",
                "🧠 Semantic similarity analysis",
                "🌍 Bilingual support (English/Arabic)",
                "📊 Comprehensive Excel reporting",
                "⚡ High-accuracy extraction"
            ],
            'upload_label': "📁 Upload CV Files",
            'upload_help': "Supported formats: PDF, DOCX, PNG, JPG, JPEG",
            'jd_label': "📋 Job Description",
            'jd_placeholder': "Paste the complete job description here...\n\nInclude required skills, qualifications, and experience.",
            'analyze_button': "🚀 Analyze CVs",
            'results_title': "📊 Analysis Results",
            'download_label': "📥 Download Comprehensive Report",
            'top_candidates': "🏆 Top Candidates",
            'summary': "📈 Summary Statistics",
            'processing': "Processing CV",
            'error_no_files': "❌ Please upload at least one CV file.",
            'error_no_jd': "❌ Please provide a job description.",
            'error_no_valid': "❌ No valid CVs could be processed. Please check file formats and content."
        },
        'ar': {
            'title': "🎯 ParsaCV - محلل السير الذاتية المتقدم متعدد اللغات",
            'subtitle': "فحص ذكي للسير الذاتية باستخدام معالجة اللغة الطبيعية المتقدمة والمطابقة الدلالية",
            'features': [
                "🤖 مطابقة أنماط معالجة اللغة الطبيعية المتقدمة",
                "🧠 تحليل التشابه الدلالي",
                "🌍 دعم ثنائي اللغة (الإنجليزية/العربية)",
                "📊 تقارير Excel شاملة",
                "⚡ استخراج عالي الدقة"
            ],
            'upload_label': "📁 رفع ملفات السيرة الذاتية",
            'upload_help': "الصيغ المدعومة: PDF, DOCX, PNG, JPG, JPEG",
            'jd_label': "📋 وصف الوظيفة",
            'jd_placeholder': "الصق وصف الوظيفة الكامل هنا...\n\nاشمل المهارات المطلوبة والمؤهلات والخبرة.",
            'analyze_button': "🚀 تحليل السير الذاتية",
            'results_title': "📊 نتائج التحليل",
            'download_label': "📥 تحميل التقرير الشامل",
            'top_candidates': "🏆 أفضل المرشحين",
            'summary': "📈 إحصائيات ملخصة",
            'processing': "معالجة السيرة الذاتية",
            'error_no_files': "❌ يرجى رفع ملف سيرة ذاتية واحد على الأقل.",
            'error_no_jd': "❌ يرجى تقديم وصف الوظيفة.",
            'error_no_valid': "❌ لم يتم معالجة أي سيرة ذاتية بنجاح. يرجى التحقق من صيغ الملفات والمحتوى."
        }
    }
    
    current_ui = ui_text[language]
    
    # Main header
    st.markdown(f"""
    <div class="main-header">
        <h1>{current_ui['title']}</h1>
        <p style="font-size: 1.2em; margin-top: 1rem;">{current_ui['subtitle']}</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Features showcase
    with st.sidebar:
        st.markdown("### ✨ Key Features")
        for feature in current_ui['features']:
            st.markdown(f"- {feature}")
    
    # Initialize analyzer
    @st.cache_resource
    def load_analyzer():
        return AdvancedCVAnalyzer()
    
    with st.spinner("🔧 Loading advanced NLP models..."):
        analyzer = load_analyzer()
    
    # Main interface
    col1, col2 = st.columns([2, 1])
    
    with col1:
        # File upload section
        st.header(current_ui['upload_label'])
        st.caption(current_ui['upload_help'])
        uploaded_files = st.file_uploader(
            "",
            type=['pdf', 'docx', 'png', 'jpg', 'jpeg'],
            accept_multiple_files=True,
            key="cv_files"
        )
        
        if uploaded_files:
            st.success(f"✅ {len(uploaded_files)} files uploaded successfully")
    
    with col2:
        if uploaded_files:
            st.info(f"📋 **Files Ready**: {len(uploaded_files)}")
            file_types = {}
            for file in uploaded_files:
                ext = Path(file.name).suffix.lower()
                file_types[ext] = file_types.get(ext, 0) + 1
            
            for ext, count in file_types.items():
                st.text(f"{ext.upper()}: {count} files")
    
    # Job Description input
    st.header(current_ui['jd_label'])
    job_description = st.text_area(
        "",
        placeholder=current_ui['jd_placeholder'],
        height=200,
        key="job_description"
    )
    
    # Analysis section
    if st.button(current_ui['analyze_button'], type="primary", use_container_width=True):
        if not uploaded_files:
            st.error(current_ui['error_no_files'])
        elif not job_description.strip():
            st.error(current_ui['error_no_jd'])
        else:
            # Create progress tracking
            progress_container = st.container()
            results_container = st.container()
            
            with progress_container:
                st.markdown("### 🔄 Processing CVs...")
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                # Process files
                results = []
                temp_dir = Path("temp_cvs")
                temp_dir.mkdir(exist_ok=True)
                
                try:
                    for i, uploaded_file in enumerate(uploaded_files):
                        # Update progress
                        progress = (i) / len(uploaded_files)
                        progress_bar.progress(progress)
                        status_text.text(f"{current_ui['processing']}: {uploaded_file.name}")
                        
                        # Save uploaded file temporarily
                        temp_file_path = temp_dir / uploaded_file.name
                        with open(temp_file_path, "wb") as f:
                            f.write(uploaded_file.read())
                        
                        # Process CV
                        result = analyzer.process_cv(str(temp_file_path), uploaded_file.name, job_description)
                        results.append(result)
                    
                    # Complete progress
                    progress_bar.progress(1.0)
                    status_text.text("✅ Processing completed!")
                    
                except Exception as e:
                    st.error(f"❌ Error during processing: {e}")
                    results = []
            
            # Display results
            with results_container:
                st.markdown("---")
                st.header(current_ui['results_title'])
                
                # Filter out None results
                valid_results = [r for r in results if r is not None]
                
                if valid_results:
                    # Summary statistics
                    st.subheader(current_ui['summary'])
                    
                    col1, col2, col3, col4 = st.columns(4)
                    
                    with col1:
                        st.metric(
                            "📊 Total CVs" if language == 'en' else "📊 إجمالي السير",
                            len(valid_results)
                        )
                    
                    with col2:
                        avg_score = sum(r['match_score'] for r in valid_results) / len(valid_results)
                        st.metric(
                            "📈 Avg Score" if language == 'en' else "📈 المتوسط",
                            f"{avg_score:.1f}%"
                        )
                    
                    with col3:
                        top_score = max(r['match_score'] for r in valid_results)
                        st.metric(
                            "🏆 Best Score" if language == 'en' else "🏆 أفضل نسبة",
                            f"{top_score:.1f}%"
                        )
                    
                    with col4:
                        high_match = len([r for r in valid_results if r['match_score'] >= 70])
                        st.metric(
                            "⭐ 70%+ Match" if language == 'en' else "⭐ +70% تطابق",
                            high_match
                        )
                    
                    # Generate and offer Excel download
                    with st.spinner("📊 Generating comprehensive report..."):
                        excel_file = analyzer.generate_excel_report(valid_results, language)
                    
                    with open(excel_file, "rb") as file:
                        st.download_button(
                            label=current_ui['download_label'],
                            data=file.read(),
                            file_name=excel_file,
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            type="primary",
                            use_container_width=True
                        )
                    
                    # Display top candidates
                    st.subheader(current_ui['top_candidates'])
                    sorted_results = sorted(valid_results, key=lambda x: x['match_score'], reverse=True)
                    
                    for i, result in enumerate(sorted_results[:5]):  # Show top 5
                        score_color = "🟢" if result['match_score'] >= 70 else "🟡" if result['match_score'] >= 50 else "🔴"
                        
                        with st.expander(f"{score_color} #{i+1} - {result['name']} - {result['match_score']:.1f}%"):
                            col1, col2 = st.columns(2)
                            
                            with col1:
                                st.markdown(f"**📧 Email:** {result['email']}")
                                st.markdown(f"**📱 Phone:** {result['phone']}")
                                st.markdown(f"**🎂 Age:** {result['age']}")
                                st.markdown(f"**🌍 Nationality:** {result['nationality']}")
                            
                            with col2:
                                st.markdown(f"**💼 Experience:** {result['total_experience_formatted']}")
                                st.markdown(f"**🗣️ Language:** {result['language'].upper()}")
                                st.markdown(f"**🎯 Confidence:** {result.get('confidence', 'medium').title()}")
                            
                            st.markdown("**✅ Matched Skills:**")
                            st.caption(result['matched_skills'][:200] + "..." if len(result['matched_skills']) > 200 else result['matched_skills'])
                            
                            st.markdown("**❌ Missing Skills:**")
                            st.caption(result['missing_skills'][:200] + "..." if len(result['missing_skills']) > 200 else result['missing_skills'])
                            
                            # Show semantic matches if available
                            if result.get('semantic_matches'):
                                st.markdown("**🧠 Best Semantic Matches:**")
                                for match in result['semantic_matches'][:3]:
                                    st.caption(f"• {match['cv_sentence']} (Score: {match['score']:.2f})")
                
                else:
                    st.error(current_ui['error_no_valid'])
                
                # Cleanup temporary files
                import shutil
                shutil.rmtree(temp_dir, ignore_errors=True)

if __name__ == "__main__":
    main()